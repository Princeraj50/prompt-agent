def classify_reports_by_domain(pdf_dir: str, domain_doc_path: str):
    """
    Classify PDF reports using a domain dictionary (.docx) and Azure OpenAI.
    
    Args:
        pdf_dir (str): Directory containing PDF files
        domain_doc_path (str): Path to Word doc with domain/subdomain definitions
    
    Returns:
        Dict[str, DomainClassification]: Mapping of filenames to structured classification output
    """

    import os
    import fitz  # PyMuPDF
    import docx
    import logging
    from typing import TypedDict, Dict

    from langchain.text_splitter import CharacterTextSplitter
    from langchain.schema import Document
    from langchain.vectorstores import FAISS
    from langchain.embeddings import AzureOpenAIEmbeddings
    from langchain.llms import AzureOpenAI
    from langchain.chains import RetrievalQA
    from langchain.prompts import PromptTemplate

    # 🧾 Define structured output types
    class DomainClassification(TypedDict):
        domain: str
        subdomain: str

    results: Dict[str, DomainClassification] = {}

    # 📝 Logger setup
    logging.basicConfig(
        filename="domain_classifier.log",
        filemode="a",
        format="%(asctime)s [%(levelname)s] %(message)s",
        level=logging.INFO
    )
    logger = logging.getLogger(__name__)
    logger.info("🚀 Starting classification pipeline")

    # 📘 Helper: Extract text and tables from Word doc
    def extract_docx_content(path: str) -> str:
        try:
            doc = docx.Document(path)
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        tables.append(" | ".join(cells))
            return "\n".join(paras + tables)
        except Exception as e:
            logger.error(f"Failed to read domain doc: {e}")
            raise RuntimeError("Could not load domain dictionary.")

    # 📄 Helper: Extract plain text from PDF
    def extract_pdf_text(path: str) -> str:
        try:
            pdf = fitz.open(path)
            return "\n".join([page.get_text() for page in pdf])
        except Exception as e:
            logger.error(f"Failed to read PDF {path}: {e}")
            raise RuntimeError(f"Could not read PDF: {path}")

    try:
        # 🔍 Load and chunk domain dictionary
        domain_text = extract_docx_content(domain_doc_path)
        logger.info(f"✅ Loaded domain dictionary from: {domain_doc_path}")

        splitter = CharacterTextSplitter(chunk_size=512, chunk_overlap=50)
        domain_chunks = splitter.split_text(domain_text)
        domain_docs = [Document(page_content=chunk) for chunk in domain_chunks]

        # 🧠 Azure Embeddings
        embedder = AzureOpenAIEmbeddings(
            model="text-embedding-ada-002",
            azure_endpoint="https://your-resource.openai.azure.com/",
            api_key="your-azure-api-key",
            api_version="2023-07-01-preview"
        )

        vectorstore = FAISS.from_documents(domain_docs, embedder)
        retriever = vectorstore.as_retriever(search_kwargs={"k": 5})
        logger.info("✅ FAISS retriever ready")

        # 🤖 Azure LLM
        llm = AzureOpenAI(
            deployment_name="your-llm-deployment-name",
            azure_endpoint="https://your-resource.openai.azure.com/",
            api_key="your-azure-api-key",
            api_version="2023-07-01-preview",
            temperature=0.1
        )
        logger.info("✅ LLM client loaded")

        # 💡 Structured Prompt
        prompt = PromptTemplate.from_template("""
You are given a client report and a domain dictionary.

Classify the report into the most relevant domain and subdomain from the dictionary.

Respond ONLY in JSON format:
{
  "domain": "<domain>",
  "subdomain": "<subdomain>"
}

Domain Dictionary:
{context}

Client Report:
{question}
""")

        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            retriever=retriever,
            chain_type="stuff",
            chain_type_kwargs={"prompt": prompt}
        )

    except Exception as setup_error:
        logger.critical(f"❌ Pipeline setup failed: {setup_error}")
        raise RuntimeError("Pipeline initialization error.")

    # 🗂️ Loop through PDFs
    for filename in os.listdir(pdf_dir):
        if filename.lower().endswith(".pdf"):
            pdf_path = os.path.join(pdf_dir, filename)
            try:
                logger.info(f"🔍 Classifying {filename}")
                report_text = extract_pdf_text(pdf_path)
                result = qa_chain.run(report_text)
                results[filename] = result
                logger.info(f"✅ {filename} → {result}")
            except Exception as e:
                logger.error(f"❌ Failed {filename}: {e}")
                results[filename] = {
                    "domain": "Unclassified",
                    "subdomain": f"Error: {str(e)}"
                }

    logger.info("🎯 Classification complete")
    return results
